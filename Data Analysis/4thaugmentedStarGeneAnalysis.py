import tarfile, os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from docx import Document

def analyze_gdc_lncrna(file_path):
    # 1. Extraction
    extract_dir = 'sample_analysis'
    os.makedirs(extract_dir, exist_ok=True)
    with tarfile.open(file_path, 'r:gz') as tar:
        tar.extractall(path=extract_dir)
        tsv_file = next(os.path.join(dp, f) for dp, dn, filenames in os.walk(extract_dir) 
                        for f in filenames if f.endswith('.tsv'))

    # 2. Load and Filter
    df = pd.read_csv(tsv_file, sep='\t', skiprows=1)
    df = df[~df['gene_id'].str.startswith('N_')].copy()
    lnc_df = df[df['gene_type'] == 'lncRNA'].copy()
    pc_df = df[df['gene_type'] == 'protein_coding'].copy()

    # 3. Top Genes
    top_10 = lnc_df.sort_values(by='tpm_unstranded', ascending=False).head(10)

    # 4. Visualization
    sns.set_theme(style="whitegrid")
    plt.figure(figsize=(10, 6))
    sns.barplot(x='tpm_unstranded', y='gene_name', data=top_10, palette='mako')
    plt.title('Top 10 Highly Expressed lncRNAs')
    plt.savefig('top_lncrnas.png')

    plt.figure(figsize=(10, 6))
    comp_df = pd.concat([lnc_df, pc_df])
    comp_df['log_tpm'] = np.log10(comp_df['tpm_unstranded'] + 0.01)
    sns.boxplot(x='gene_type', y='log_tpm', data=comp_df)
    plt.title('Expression Distribution Comparison')
    plt.savefig('expression_distribution.png')

    # 5. Export Report
    doc = Document()
    doc.add_heading('lncRNA Analysis Report', 0)
    table = doc.add_table(rows=1, cols=2)
    for _, row in top_10.head(5).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['gene_name']
        row_cells[1].text = f"{row['tpm_unstranded']:.2f} TPM"
    doc.save('lncRNA_Analysis_Report.docx')
    print("Analysis complete. Files generated: top_lncrnas.png, expression_distribution.png, lncRNA_Analysis_Report.docx")

# Run the analysis
analyze_gdc_lncrna('gdc_download_20260304_201550.898969.tar.gz')
